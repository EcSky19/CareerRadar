"""
Resume Parser
=============
Parses uploaded resumes (PDF, DOCX, TXT) into:
  1. Raw plain text
  2. Structured JSON sections:
       { contact, education, experience, projects, skills,
         certifications, leadership, awards }

Design notes:
  - Uses pdfminer.six for PDF (no external binaries needed)
  - Uses python-docx for DOCX
  - Extracts sections using heading / keyword heuristics
  - Does NOT use AI for parsing — purely text processing
    (AI analysis happens in analyzer.py)
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)

# Section heading keywords → canonical section names
_SECTION_PATTERNS: list[tuple[re.Pattern, str]] = [
    (re.compile(r"(experience|work history|employment)", re.I), "experience"),
    (re.compile(r"(education|academic|university|college)", re.I), "education"),
    (re.compile(r"(project|projects|personal project)", re.I), "projects"),
    (re.compile(r"(skills|technical skills|technologies|stack)", re.I), "skills"),
    (re.compile(r"(certif)", re.I), "certifications"),
    (re.compile(r"(leadership|activities|clubs|organizations)", re.I), "leadership"),
    (re.compile(r"(award|honor|achievement)", re.I), "awards"),
    (re.compile(r"(summary|objective|about|profile)", re.I), "summary"),
]


@dataclass
class ParsedResume:
    raw_text: str = ""
    contact: dict           = field(default_factory=dict)
    education: list[dict]   = field(default_factory=list)
    experience: list[dict]  = field(default_factory=list)
    projects: list[dict]    = field(default_factory=list)
    skills: list[str]       = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    leadership: list[str]   = field(default_factory=list)
    awards: list[str]       = field(default_factory=list)
    summary: str            = ""
    parse_warnings: list[str] = field(default_factory=list)

    def to_json(self) -> dict:
        return {
            "contact":        self.contact,
            "education":      self.education,
            "experience":     self.experience,
            "projects":       self.projects,
            "skills":         self.skills,
            "certifications": self.certifications,
            "leadership":     self.leadership,
            "awards":         self.awards,
            "summary":        self.summary,
        }


# =============================================================================
# PUBLIC API
# =============================================================================

async def parse_resume(content: bytes, filename: str) -> ParsedResume:
    """
    Parse a resume file and return a ParsedResume.
    Detects format from filename extension.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _parse_pdf(content)
    elif ext in ("docx", "doc"):
        return _parse_docx(content)
    elif ext in ("txt", "text", "md"):
        return _parse_text(content.decode("utf-8", errors="replace"))
    else:
        # Try PDF first, then plain text
        try:
            return _parse_pdf(content)
        except Exception:
            return _parse_text(content.decode("utf-8", errors="replace"))


# =============================================================================
# FORMAT-SPECIFIC EXTRACTORS
# =============================================================================

def _parse_pdf(content: bytes) -> ParsedResume:
    try:
        from pdfminer.high_level import extract_text
    except ImportError:
        raise ImportError("pdfminer.six is required: pip install pdfminer.six")

    try:
        text = extract_text(io.BytesIO(content))
    except Exception as exc:
        result = ParsedResume()
        result.parse_warnings.append(f"PDF extraction failed: {exc}")
        return result

    return _parse_text(text)


def _parse_docx(content: bytes) -> ParsedResume:
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    try:
        doc = docx.Document(io.BytesIO(content))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        # Also extract from tables (common in resume templates)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in lines:
                        lines.append(cell_text)
        full_text = "\n".join(lines)
    except Exception as exc:
        result = ParsedResume()
        result.parse_warnings.append(f"DOCX extraction failed: {exc}")
        return result

    return _parse_text(full_text)


def _parse_text(text: str) -> ParsedResume:
    result = ParsedResume(raw_text=text)
    lines  = [l.rstrip() for l in text.splitlines()]

    # Extract contact info from top 10 lines
    result.contact = _extract_contact("\n".join(lines[:10]))

    # Split into sections
    sections = _split_into_sections(lines)

    for section_name, section_lines in sections.items():
        section_text = "\n".join(section_lines).strip()

        if section_name == "experience":
            result.experience = _parse_experience_section(section_lines)
        elif section_name == "education":
            result.education = _parse_education_section(section_lines)
        elif section_name == "projects":
            result.projects = _parse_projects_section(section_lines)
        elif section_name == "skills":
            result.skills = _parse_skills_section(section_lines)
        elif section_name == "certifications":
            result.certifications = [l.strip("•- ") for l in section_lines if l.strip()]
        elif section_name == "leadership":
            result.leadership = [l.strip("•- ") for l in section_lines if l.strip()]
        elif section_name == "awards":
            result.awards = [l.strip("•- ") for l in section_lines if l.strip()]
        elif section_name == "summary":
            result.summary = section_text

    if not result.experience and not result.projects and not result.skills:
        result.parse_warnings.append(
            "Could not detect standard resume sections. "
            "Raw text is preserved for analysis."
        )

    return result


# =============================================================================
# SECTION SPLITTER
# =============================================================================

def _split_into_sections(lines: list[str]) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current_section = "header"
    sections[current_section] = []

    for line in lines:
        matched_section = _detect_section_heading(line)
        if matched_section:
            current_section = matched_section
            if current_section not in sections:
                sections[current_section] = []
        else:
            sections[current_section].append(line)

    return sections


def _detect_section_heading(line: str) -> Optional[str]:
    """
    Returns the canonical section name if the line appears to be a section header,
    otherwise returns None.

    Heuristics:
    - Short (< 50 chars)
    - Matches a section keyword
    - No bullet, no indentation, no comma (avoids matching job titles)
    """
    stripped = line.strip()
    if not stripped or len(stripped) > 50:
        return None
    if stripped.startswith(("•", "-", "*", "·")):
        return None
    if "," in stripped and len(stripped.split()) > 3:
        return None

    for pattern, name in _SECTION_PATTERNS:
        if pattern.search(stripped):
            return name

    return None


# =============================================================================
# SECTION PARSERS
# =============================================================================

def _extract_contact(text: str) -> dict:
    contact: dict = {}

    # Email
    email = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    if email:
        contact["email"] = email.group(0)

    # Phone
    phone = re.search(r"(\+?1?\s?)?(\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4})", text)
    if phone:
        contact["phone"] = phone.group(0).strip()

    # LinkedIn
    linkedin = re.search(r"linkedin\.com/in/([^\s,/\"']+)", text, re.I)
    if linkedin:
        contact["linkedin"] = f"linkedin.com/in/{linkedin.group(1)}"

    # GitHub
    github = re.search(r"github\.com/([^\s,/\"']+)", text, re.I)
    if github:
        contact["github"] = f"github.com/{github.group(1)}"

    # First line is often the name
    first_line = text.strip().splitlines()[0].strip() if text.strip() else ""
    if first_line and len(first_line) < 60 and not re.search(r"[@/]", first_line):
        contact["name"] = first_line

    return contact


def _parse_experience_section(lines: list[str]) -> list[dict]:
    entries = []
    current: Optional[dict] = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Detect new entry: lines that look like "Company | Title | Date" or
        # a standalone company name followed by indented bullets
        if _looks_like_entry_header(stripped):
            if current:
                entries.append(current)
            current = {
                "header": stripped,
                "bullets": [],
                "company": _extract_company(stripped),
                "title": _extract_title_from_header(stripped),
                "dates": _extract_dates(stripped),
            }
        elif current and stripped.startswith(("•", "-", "*", "·")):
            current["bullets"].append(stripped.lstrip("•-*· "))
        elif current:
            # Might be a continuation line or sub-title
            if not current.get("title") and len(stripped) < 80:
                current["title"] = stripped

    if current:
        entries.append(current)

    return entries


def _parse_education_section(lines: list[str]) -> list[dict]:
    entries = []
    current: Optional[dict] = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if _looks_like_entry_header(stripped) or re.search(r"\b(university|college|school|institute|b\.s\.|b\.a\.|m\.s\.|phd)\b", stripped, re.I):
            if current:
                entries.append(current)
            current = {
                "header": stripped,
                "details": [],
                "dates": _extract_dates(stripped),
            }
        elif current:
            current["details"].append(stripped)

    if current:
        entries.append(current)

    return entries


def _parse_projects_section(lines: list[str]) -> list[dict]:
    entries = []
    current: Optional[dict] = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if _looks_like_entry_header(stripped) and not stripped.startswith(("•", "-")):
            if current:
                entries.append(current)
            current = {"name": stripped, "bullets": []}
        elif current:
            current["bullets"].append(stripped.lstrip("•-*· "))

    if current:
        entries.append(current)

    return entries


def _parse_skills_section(lines: list[str]) -> list[str]:
    skills = []
    for line in lines:
        # Split on commas, pipes, semicolons, bullets
        parts = re.split(r"[,|;•·\n]", line)
        for part in parts:
            s = part.strip().strip("-·• ")
            if s and 1 < len(s) < 60:
                skills.append(s)
    return list(dict.fromkeys(skills))   # deduplicate, preserve order


# =============================================================================
# HELPERS
# =============================================================================

def _looks_like_entry_header(line: str) -> bool:
    """Lines that look like a job/school entry rather than a bullet."""
    if line.startswith(("•", "-", "*")):
        return False
    if len(line) > 120:
        return False
    # Has a year range or company-like formatting
    if re.search(r"\b(20\d{2}|19\d{2}|present|current)\b", line, re.I):
        return True
    # Title Case with limited length (likely a heading)
    words = line.split()
    if 1 < len(words) <= 8 and sum(1 for w in words if w and w[0].isupper()) >= len(words) * 0.6:
        return True
    return False


def _extract_dates(text: str) -> str:
    m = re.search(
        r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|\d{4})"
        r".*(present|current|\d{4})",
        text, re.I
    )
    return m.group(0) if m else ""


def _extract_company(text: str) -> str:
    # Very simple: take text before first " | " or " - " if present
    for sep in (" | ", " – ", " - ", " · "):
        if sep in text:
            return text.split(sep)[0].strip()
    return text.strip()


def _extract_title_from_header(text: str) -> str:
    for sep in (" | ", " – ", " - ", " · "):
        if sep in text:
            parts = text.split(sep)
            if len(parts) >= 2:
                return parts[1].strip()
    return ""
